# -*- coding: utf-8 -*-
from datetime import time

import pandas as pd
################################################################################
## Form generated from reading UI file 'spravochnik_investments.ui'
##
## Created by: Qt User Interface Compiler version 6.7.2
##
## WARNING! All changes made in this file will be lost when recompiling UI file!
################################################################################

from PySide6.QtCore import (QCoreApplication, QDate, QDateTime, QLocale,
    QMetaObject, QObject, QPoint, QRect,
    QSize, QTime, QUrl, Qt, Slot)
from PySide6.QtGui import (QBrush, QColor, QConicalGradient, QCursor,
                           QFont, QFontDatabase, QGradient, QIcon,
                           QImage, QKeySequence, QLinearGradient, QPainter,
                           QPalette, QPixmap, QRadialGradient, QTransform, QShortcut)
from PySide6.QtSql import QSqlQuery, QSqlDatabase
from PySide6.QtWidgets import (QApplication, QHeaderView, QLabel, QLineEdit,
                               QPushButton, QSizePolicy, QTableWidget, QTableWidgetItem,
                               QWidget, QDialog)
from sqlalchemy import create_engine, text

from settings import DB_PATH
from ui_compiled import investment_add, investment_edit
from time import time
import docx


class Ui_Form(object):
    data_for_table_global=pd.DataFrame()
    def search_investment(self):

        TABLE_ROW_LIMIT = 10

        if self.data_for_table_global.empty:
            db_connection = create_engine(f'sqlite:///{DB_PATH}').connect()

            data_for_table = pd.read_sql(text(f'SELECT investments.pk, customers.name, papers.info,'
                                          f'investments.paper_pk,investments.customer_pk, investments.kotirovka, investments.buy_date,investments.sell_date '
                                          f'FROM investments '
                                          f'JOIN customers on customers.pk=investments.customer_pk '
                                          f'JOIN papers on papers.pk=investments.paper_pk'), db_connection).astype(str)
            self.tableWidget.setRowCount(len(data_for_table))
            self.tableWidget.setColumnCount(len(data_for_table.columns))


            hide_musor = True
            if hide_musor:
                self.tableWidget.setColumnHidden(3, True)
                self.tableWidget.setColumnHidden(4, True)
                self.tableWidget.setColumnHidden(5, True)
                self.tableWidget.setColumnHidden(6, True)
                self.tableWidget.setColumnHidden(7, True)
        else:
            data_for_table=self.data_for_table_global
#========================================================
        string_for_search = self.searchLineEdit.text()
        column_for_search = self.tableWidget.currentColumn()
        print(f"string {string_for_search}, column {column_for_search}")
        df = data_for_table

        df=df[df[df.columns[column_for_search]] == string_for_search]

        data_for_table=df

        self.tableWidget.setRowCount(0)
        self.tableWidget.setColumnCount(0)

        self.tableWidget.setRowCount(len(data_for_table))
        self.tableWidget.setColumnCount(len(data_for_table.columns))
        hide_musor = True
        if hide_musor:
            self.tableWidget.setColumnHidden(3, True)
            self.tableWidget.setColumnHidden(4, True)
            self.tableWidget.setColumnHidden(5, True)
            self.tableWidget.setColumnHidden(6, True)
            self.tableWidget.setColumnHidden(7, True)

        print(data_for_table)
        self.data_for_table_global=data_for_table
        for col_num in range(len(data_for_table.columns)):
            for row_num in range(len(data_for_table)):
                self.tableWidget.setItem(row_num, col_num,
                                         QTableWidgetItem(data_for_table.iloc[row_num, col_num]))

    def delete_investment(self):
        VetDbConnnection = QSqlDatabase.addDatabase("QSQLITE")
        VetDbConnnection.setDatabaseName(DB_PATH)
        VetDbConnnection.open()
        VetTableQuery = QSqlQuery()
        vet_query_str = """DELETE FROM investments WHERE pk =:pk"""
        VetTableQuery.prepare(vet_query_str)

        VetTableQuery.bindValue(":pk", self.tableWidget.item(self.tableWidget.currentRow(),0).text())

        uspeh = VetTableQuery.exec()
        VetDbConnnection.close()

        self.fill_investments_table()

        print("uspeh", uspeh)
    def open_investments_add(self):
        self.window = QDialog()
        self.ui = investment_add.Ui_Form()
        self.ui.setupUi(self.window)
        self.window.show()
    def open_investments_edit(self):
        self.window = QDialog()
        self.ui = investment_edit.Ui_Form()
        self.ui.transfer_data(self.tableWidget.item(self.tableWidget.currentRow(),4).text(),
                              self.tableWidget.item(self.tableWidget.currentRow(),3).text(),
                              self.tableWidget.item(self.tableWidget.currentRow(),5).text(),
                              self.tableWidget.item(self.tableWidget.currentRow(),6).text(),
                              self.tableWidget.item(self.tableWidget.currentRow(),7).text(),
                              self.tableWidget.item(self.tableWidget.currentRow(),0).text(),)
        self.ui.setupUi(self.window)
        self.window.show()

    def print_document(self):
        import docx
        import pandas as pd

        # i am not sure how you are getting your data, but you said it is a
        # pandas data frame
        if self.data_for_table_global.empty:
            db_connection = create_engine(f'sqlite:///{DB_PATH}').connect()

            data_for_table = pd.read_sql(text(f'SELECT investments.pk, customers.name, papers.info,'
                                              f'investments.paper_pk,investments.customer_pk, investments.kotirovka, investments.buy_date,investments.sell_date '
                                              f'FROM investments '
                                              f'JOIN customers on customers.pk=investments.customer_pk '
                                              f'JOIN papers on papers.pk=investments.paper_pk'), db_connection).astype(
                str)
            self.tableWidget.setRowCount(len(data_for_table))
            self.tableWidget.setColumnCount(len(data_for_table.columns))

            hide_musor = True
            if hide_musor:
                self.tableWidget.setColumnHidden(3, True)
                self.tableWidget.setColumnHidden(4, True)
                self.tableWidget.setColumnHidden(5, True)
                self.tableWidget.setColumnHidden(6, True)
                self.tableWidget.setColumnHidden(7, True)
        else:
            data_for_table = self.data_for_table_global
        # ========================================================        if self.data_for_table_global.empty:
        #             db_connection = create_engine(f'sqlite:///{DB_PATH}').connect()
        #
        #             data_for_table = pd.read_sql(text(f'SELECT investments.pk, customers.name, papers.info,'
        #                                           f'investments.paper_pk,investments.customer_pk, investments.kotirovka, investments.buy_date,investments.sell_date '
        #                                           f'FROM investments '
        #                                           f'JOIN customers on customers.pk=investments.customer_pk '
        #                                           f'JOIN papers on papers.pk=investments.paper_pk'), db_connection).astype(str)
        #             self.tableWidget.setRowCount(len(data_for_table))
        #             self.tableWidget.setColumnCount(len(data_for_table.columns))
        #
        #
        #             hide_musor = True
        #             if hide_musor:
        #                 self.tableWidget.setColumnHidden(3, True)
        #                 self.tableWidget.setColumnHidden(4, True)
        #                 self.tableWidget.setColumnHidden(5, True)
        #                 self.tableWidget.setColumnHidden(6, True)
        #                 self.tableWidget.setColumnHidden(7, True)
        #         else:
        #             data_for_table=self.data_for_table_global
        # #========================================================
        df=data_for_table
        # open an existing document
        doc = docx.Document()
        doc.save('test.docx')
        doc.add_paragraph("Отчёт о скуфах"
                          "")

        # add a table to the end and create a reference variable
        # extra row is so we can add the header row
        t = doc.add_table(df.shape[0] + 1, df.shape[1])

        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0, j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i + 1, j).text = str(df.values[i, j])

        # save the doc
        doc.save(f'C:/MHAD/Отчёт о скуфах{int(time())}.docx')
    @Slot()
    def fill_investments_table(self):
        TABLE_ROW_LIMIT = 10
        self.data_for_table_global=pd.DataFrame()
        db_connection = create_engine(f'sqlite:///{DB_PATH}').connect()

        data_for_table = pd.read_sql(text(f'SELECT investments.pk, customers.name, papers.info,'
                                          f'investments.paper_pk,investments.customer_pk, investments.kotirovka, investments.buy_date,investments.sell_date '
                                          f'FROM investments '
                                          f'JOIN customers on customers.pk=investments.customer_pk '
                                          f'JOIN papers on papers.pk=investments.paper_pk'), db_connection).astype(str)
        self.tableWidget.setRowCount(len(data_for_table))
        self.tableWidget.setColumnCount(len(data_for_table.columns))
        hide_musor = True
        if hide_musor:
            self.tableWidget.setColumnHidden(3,True)
            self.tableWidget.setColumnHidden(4,True)
            self.tableWidget.setColumnHidden(5,True)
            self.tableWidget.setColumnHidden(6,True)
            self.tableWidget.setColumnHidden(7,True)
        #print(data_for_table,len(data_for_table),(len(data_for_table.columns)))
        for col_num in range(len(data_for_table.columns)):
            for row_num in range(len(data_for_table)):
                self.tableWidget.setItem(row_num, col_num,
                                                 QTableWidgetItem(data_for_table.iloc[row_num, col_num]))
    def setupUi(self, Form):


        if not Form.objectName():
            Form.setObjectName(u"Form")
        Form.resize(724, 507)
        self.addPushButton = QPushButton(Form, clicked = lambda:self.open_investments_add())
        self.addPushButton.setObjectName(u"addPushButton")
        self.addPushButton.setGeometry(QRect(40, 420, 111, 61))
        self.tableWidget = QTableWidget(Form)
        self.tableWidget.setObjectName(u"tableWidget")
        self.tableWidget.setGeometry(QRect(30, 110, 651, 291))
        self.editPushButton = QPushButton(Form, clicked = lambda :self.open_investments_edit())
        self.editPushButton.setObjectName(u"editPushButton")
        self.editPushButton.setGeometry(QRect(180, 420, 111, 61))
        self.deletePushButton = QPushButton(Form, clicked = lambda:self.delete_investment())
        self.deletePushButton.setObjectName(u"deletePushButton")
        self.deletePushButton.setGeometry(QRect(580, 420, 111, 61))
        self.searchPushButton = QPushButton(Form, clicked = lambda:self.search_investment())
        self.searchPushButton.setObjectName(u"searchPushButton")
        self.searchPushButton.setGeometry(QRect(310, 60, 111, 31))
        self.label = QLabel(Form)
        self.label.setObjectName(u"label")
        self.label.setGeometry(QRect(40, 20, 281, 31))
        font = QFont()
        font.setPointSize(19)
        self.label.setFont(font)
        self.searchLineEdit = QLineEdit(Form)
        self.searchLineEdit.setObjectName(u"searchLineEdit")
        self.searchLineEdit.setGeometry(QRect(40, 70, 251, 21))

        self.documentPushButton = QPushButton(Form, clicked = lambda:self.print_document())
        self.documentPushButton.setObjectName(u"documentPushButton")
        self.documentPushButton.setGeometry(QRect(310, 420, 141, 61))
        self.documentPushButton.setText("Выпуск документа")

        self.retranslateUi(Form)

        self.fill_investments_table()

        self.shortcut = QShortcut(QKeySequence("Ctrl+R"), Form)
        self.shortcut.activated.connect(lambda: self.fill_investments_table())

        QMetaObject.connectSlotsByName(Form)
    # setupUi

    def retranslateUi(self, Form):
        Form.setWindowTitle(QCoreApplication.translate("Form", u"Form", None))
        self.addPushButton.setText(QCoreApplication.translate("Form", u"\u0414\u043e\u0431\u0430\u0432\u0438\u0442\u044c", None))
        self.editPushButton.setText(QCoreApplication.translate("Form", u"\u0440\u0435\u0434\u0430\u043a\u0442\u0438\u0440\u043e\u0432\u0430\u0442\u044c", None))
        self.deletePushButton.setText(QCoreApplication.translate("Form", u"\u0423\u0434\u0430\u043b\u0438\u0442\u044c", None))
        self.searchPushButton.setText(QCoreApplication.translate("Form", u"\u0438\u0441\u043a\u0430\u0442\u044c", None))
        self.label.setText(QCoreApplication.translate("Form", u"\u0421\u043f\u0440\u0430\u0432\u043e\u0447\u043d\u0438\u043a \u0418\u043d\u0432\u0435\u0441\u0442\u0438\u0446\u0438\u0439", None))
        self.documentPushButton.setText("Документ")
    # retranslateUi

